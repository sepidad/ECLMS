"""DOCX and PDF representations of a structured contract version."""

from __future__ import annotations

from io import BytesIO
from typing import Any
from html.parser import HTMLParser
from base64 import b64decode

from backend.modules.contracts.domain.structure import numbered_structure
from backend.modules.contracts.domain.rich_text import rich_text_to_plain


STYLE_PROFILE = {
  'body_font': 'Aptos',
  'title_font': 'Aptos Display',
  'body_size': 10.5,
  'article_size': 13,
  'sub_article_size': 11,
  'accent': '1F4E79',
  'note_color': '666666',
}


def _active(version_items: list[dict[str, Any]]) -> dict[str, Any]:
  return next((item for item in version_items if item.get('is_active')), version_items[-1] if version_items else {})


def _structure(version: dict[str, Any]) -> list[dict[str, Any]]:
  structure = version.get('structure')
  if structure:
    return structure
  content = (version.get('content') or '').strip()
  return [{'id': 'legacy', 'title': 'Contract content', 'body': content, 'children': [], 'notes': []}] if content else []


def _add_page_field(paragraph) -> None:
  from docx.oxml import OxmlElement
  from docx.oxml.ns import qn
  run = paragraph.add_run()
  begin = OxmlElement('w:fldChar'); begin.set(qn('w:fldCharType'), 'begin')
  instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = ' PAGE '
  end = OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'), 'end')
  run._r.extend([begin, instr, end])


class _DocxRichParser(HTMLParser):
  """Parse the editor's small HTML subset into paragraphs and tables."""
  def __init__(self) -> None:
    super().__init__(convert_charrefs=True)
    self.blocks: list[tuple[str, Any]] = []
    self.tokens: list[tuple[str, str, dict[str, str]]] = []
    self.stack: list[str] = []
    self.block_style: dict[str, str] = {}
    self.list_kind: str | None = None
    self.table: list[list[str]] | None = None
    self.row: list[str] | None = None
    self.cell: list[str] | None = None

  def _flush(self) -> None:
    if self.tokens:
      self.blocks.append(('paragraph', (self.tokens, self.block_style, self.list_kind)))
      self.tokens = []
    self.block_style = {}

  def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
    attributes = {key: value or '' for key, value in attrs}
    if tag in {'p', 'div', 'li'}:
      self._flush()
      self.block_style = {'style': attributes.get('style', ''), 'dir': attributes.get('dir', '')}
      if tag == 'li': self.list_kind = self.list_kind or 'ul'
    elif tag in {'ul', 'ol'}:
      self.list_kind = tag
    elif tag == 'table':
      self._flush(); self.table = []
    elif tag == 'tr' and self.table is not None:
      self.row = []
    elif tag in {'td', 'th'} and self.row is not None:
      self.cell = []
    elif tag in {'strong', 'b', 'em', 'i', 'u'}:
      self.stack.append(tag)
    elif tag == 'br':
      self.tokens.append(('text', '\n', {'bold': '1' if any(x in {'strong', 'b'} for x in self.stack) else '', 'italic': '1' if any(x in {'em', 'i'} for x in self.stack) else '', 'underline': '1' if 'u' in self.stack else ''}))
    elif tag == 'img':
      self.tokens.append(('image', attributes.get('src', ''), {}))

  def handle_endtag(self, tag: str) -> None:
    if tag in {'p', 'div', 'li'}: self._flush()
    elif tag in {'ul', 'ol'}: self.list_kind = None
    elif tag in {'strong', 'b', 'em', 'i', 'u'} and tag in self.stack: self.stack.remove(tag)
    elif tag in {'td', 'th'} and self.row is not None and self.cell is not None:
      self.row.append(''.join(self.cell).strip()); self.cell = None
    elif tag == 'tr' and self.table is not None and self.row is not None:
      self.table.append(self.row); self.row = None
    elif tag == 'table' and self.table is not None:
      self.blocks.append(('table', self.table)); self.table = None

  def handle_data(self, data: str) -> None:
    if self.cell is not None: self.cell.append(data)
    elif self.table is None: self.tokens.append(('text', data, {'bold': '1' if any(x in {'strong', 'b'} for x in self.stack) else '', 'italic': '1' if any(x in {'em', 'i'} for x in self.stack) else '', 'underline': '1' if 'u' in self.stack else ''}))


def add_rich_docx_content(document, html: str) -> None:
  from docx.enum.text import WD_ALIGN_PARAGRAPH
  from docx.shared import Inches
  parser = _DocxRichParser(); parser.feed(html); parser.close(); parser._flush()
  for kind, payload in parser.blocks:
    if kind == 'table':
      rows = payload
      if not rows: continue
      table = document.add_table(rows=len(rows), cols=max(len(row) for row in rows))
      table.style = 'Table Grid'
      for row_index, row in enumerate(rows):
        for cell_index, text in enumerate(row): table.cell(row_index, cell_index).text = text
      continue
    tokens, block_style, list_kind = payload
    style = 'List Bullet' if list_kind == 'ul' else 'List Number' if list_kind == 'ol' else None
    paragraph = document.add_paragraph(style=style)
    css = block_style.get('style', '')
    if 'text-align:center' in css: paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif 'text-align:right' in css: paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif 'text-align:justify' in css: paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for token_kind, value, flags in tokens:
      if token_kind == 'image' and value.startswith('data:image/') and ',' in value:
        try: paragraph.add_run().add_picture(BytesIO(b64decode(value.split(',', 1)[1])), width=Inches(5.8))
        except Exception: pass
      elif token_kind == 'text':
        run = paragraph.add_run(value); run.bold = bool(flags.get('bold')); run.italic = bool(flags.get('italic')); run.underline = bool(flags.get('underline'))


def build_docx(*, title: str, reference: str, counterparty: str, version: dict[str, Any], template_bytes: bytes | None = None) -> bytes:
  from io import BytesIO
  from docx import Document
  from docx.enum.section import WD_SECTION
  from docx.enum.text import WD_ALIGN_PARAGRAPH
  from docx.enum.style import WD_STYLE_TYPE
  from docx.shared import Inches, Pt, RGBColor

  document = Document(BytesIO(template_bytes)) if template_bytes else Document()
  section = document.sections[0]
  if not template_bytes:
    section.top_margin = Inches(0.8); section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9); section.right_margin = Inches(0.9)
  styles = document.styles
  def ensure_style(name: str, base: str | None = None):
    try:
      return styles[name]
    except KeyError:
      style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
      if base:
        style.base_style = ensure_style(base)
      return style

  normal = ensure_style('Normal'); normal.font.name = STYLE_PROFILE['body_font']; normal.font.size = Pt(STYLE_PROFILE['body_size'])
  for style_name, size, color in [('Title', 20, STYLE_PROFILE['accent']), ('Heading 1', STYLE_PROFILE['article_size'], STYLE_PROFILE['accent']), ('Heading 2', STYLE_PROFILE['sub_article_size'], STYLE_PROFILE['accent'])]:
    style = ensure_style(style_name, 'Normal'); style.font.name = STYLE_PROFILE['title_font']; style.font.size = Pt(size); style.font.bold = True; style.font.color.rgb = RGBColor.from_string(color)
  bullet_style = ensure_style('List Bullet', 'Normal'); bullet_style.font.name = STYLE_PROFILE['body_font']; bullet_style.font.size = Pt(STYLE_PROFILE['body_size'])

  header = section.header.paragraphs[0]
  if not template_bytes:
    header.text = 'ECLMS | CONTRACT DOCUMENT'
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.name = STYLE_PROFILE['body_font']; header.runs[0].font.size = Pt(8); header.runs[0].font.color.rgb = RGBColor(100, 116, 139)
  footer = section.footer.paragraphs[0]
  if not template_bytes or not footer.text.strip():
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run('Page '); _add_page_field(footer)

  heading = document.add_paragraph(style='Title'); heading.alignment = WD_ALIGN_PARAGRAPH.CENTER; heading.add_run(title)
  meta = document.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
  meta.add_run(f'{reference}  |  {counterparty}').italic = True
  document.add_paragraph()

  numbered, _, _ = numbered_structure(_structure(version))
  def visit(nodes: list[dict[str, Any]], depth: int = 0) -> None:
    for node in nodes:
      style = 'Heading 1' if depth == 0 else 'Heading 2'
      p = document.add_paragraph(style=style)
      p.add_run(f"Article {node['number']} - {node['title']}" if depth == 0 else f"{node['number']} {node['title']}")
      if node.get('body'):
        add_rich_docx_content(document, str(node['body']))
      for note in node.get('notes', []):
        add_rich_docx_content(document, f'<ul><li>{note}</li></ul>')
      visit(node.get('children', []), depth + 1)
  visit(numbered)
  out = BytesIO(); document.save(out); return out.getvalue()


def build_pdf(*, title: str, reference: str, counterparty: str, version: dict[str, Any], template_bytes: bytes | None = None) -> bytes:
  from reportlab.lib import colors
  from reportlab.lib.enums import TA_CENTER, TA_RIGHT
  from reportlab.lib.pagesizes import A4
  from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
  from reportlab.lib.units import inch
  from reportlab.platypus import BaseDocTemplate, Frame, PageTemplate, Paragraph, PageBreak, Spacer

  header_image = None
  if template_bytes:
    from zipfile import ZipFile
    try:
      with ZipFile(BytesIO(template_bytes)) as archive:
        image_names = sorted(name for name in archive.namelist() if name.startswith('word/media/'))
        if image_names:
          from reportlab.lib.utils import ImageReader
          header_image = ImageReader(BytesIO(archive.read(image_names[0])))
    except Exception:  # noqa: BLE001 - a malformed optional header must not block PDF export
      header_image = None

  out = BytesIO()
  styles = getSampleStyleSheet()
  body = ParagraphStyle('ContractBody', parent=styles['BodyText'], fontName='Helvetica', fontSize=10, leading=14, spaceAfter=7)
  title_style = ParagraphStyle('ContractTitle', parent=styles['Title'], fontName='Helvetica-Bold', fontSize=19, leading=23, alignment=TA_CENTER, textColor=colors.HexColor('#1F4E79'), spaceAfter=8)
  article = ParagraphStyle('Article', parent=styles['Heading1'], fontName='Helvetica-Bold', fontSize=13, leading=16, textColor=colors.HexColor('#1F4E79'), spaceBefore=10, spaceAfter=5)
  sub = ParagraphStyle('SubArticle', parent=styles['Heading2'], fontName='Helvetica-Bold', fontSize=11, leading=14, textColor=colors.HexColor('#1F4E79'), spaceBefore=7, spaceAfter=4)
  note = ParagraphStyle('Note', parent=body, leftIndent=18, bulletIndent=7, textColor=colors.HexColor('#666666'))

  def furniture(canvas, doc):
    canvas.saveState(); canvas.setFillColor(colors.HexColor('#64748B'))
    if header_image:
      canvas.drawImage(header_image, 0.65 * inch, A4[1] - 1.25 * inch, width=A4[0] - 1.3 * inch, height=1.0 * inch, preserveAspectRatio=True, anchor='n', mask='auto')
    else:
      canvas.setFont('Helvetica', 8)
      canvas.drawRightString(A4[0] - 0.9 * inch, A4[1] - 0.48 * inch, 'ECLMS | CONTRACT DOCUMENT')
    canvas.setFont('Helvetica', 8)
    canvas.drawCentredString(A4[0] / 2, 0.42 * inch, f'Page {doc.page}')
    canvas.restoreState()

  doc = BaseDocTemplate(out, pagesize=A4, leftMargin=0.9 * inch, rightMargin=0.9 * inch, topMargin=1.45 * inch if header_image else 0.75 * inch, bottomMargin=0.65 * inch)
  doc.addPageTemplates([PageTemplate(id='contract', frames=[Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id='normal')], onPage=furniture)])
  story = [Paragraph(title, title_style), Paragraph(f'{reference} | {counterparty}', ParagraphStyle('Meta', parent=body, alignment=TA_CENTER, textColor=colors.HexColor('#64748B'))), Spacer(1, 14)]
  numbered, _, _ = numbered_structure(_structure(version))
  def visit(nodes: list[dict[str, Any]], depth: int = 0):
    for node in nodes:
      story.append(Paragraph(f"Article {node['number']} - {node['title']}" if depth == 0 else f"{node['number']} {node['title']}", article if depth == 0 else sub))
      for paragraph in str(node.get('body') or '').split('\n'):
        plain = rich_text_to_plain(paragraph)
        if plain.strip(): story.append(Paragraph(plain.replace('&', '&amp;'), body))
      for note_text in node.get('notes', []): story.append(Paragraph(rich_text_to_plain(note_text).replace('&', '&amp;'), note, bulletText='•'))
      visit(node.get('children', []), depth + 1)
  visit(numbered); doc.build(story); return out.getvalue()
